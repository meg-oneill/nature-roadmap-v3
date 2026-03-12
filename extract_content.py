#!/usr/bin/env python3
"""Extract ALL text content from index.html into a structured Word document.
V3: Comprehensive extraction including timeline, challenge/response details,
    stakeholder tabs, panel data, actions, stats, tags, and all hidden content."""

from bs4 import BeautifulSoup, NavigableString, Comment
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Read HTML
with open('/home/user/nature-roadmap-v3/index.html', 'r', encoding='utf-8') as f:
    html = f.read()

# ── Extract JavaScript data BEFORE parsing ──

# Extract panelData
panel_data = {}
panel_match = re.search(r'const panelData\s*=\s*\{(.+?)\n\};', html, re.DOTALL)
if panel_match:
    raw = panel_match.group(0)
    panel_pattern = re.compile(
        r"'([^']+)':\s*\{[^}]*"
        r"principle:\s*'([^']*)'[^}]*"
        r"title:\s*'([^']*)'[^}]*"
        r"time:\s*'([^']*)'[^}]*"
        r"context:\s*'((?:[^'\\]|\\.)*)'[^}]*"
        r"actions:\s*\[(.*?)\][^}]*"
        r"gbca:\s*\[(.*?)\][^}]*"
        r"industry:\s*\[(.*?)\]",
        re.DOTALL
    )
    def extract_items(s):
        return [x.strip().replace("\\'", "'") for x in re.findall(r"'((?:[^'\\]|\\.)*)'", s)]
    for m in panel_pattern.finditer(raw):
        panel_data[m.group(1)] = {
            'principle': m.group(2), 'title': m.group(3), 'time': m.group(4),
            'context': m.group(5).replace("\\'", "'"),
            'actions': extract_items(m.group(6)),
            'gbca': extract_items(m.group(7)),
            'industry': extract_items(m.group(8)),
        }
    print(f"Extracted {len(panel_data)} panels from panelData")

# Extract ACTIONS array
actions_data = []
actions_match = re.search(r'const ACTIONS\s*=\s*\[(.*?)\];', html, re.DOTALL)
if actions_match:
    action_pattern = re.compile(
        r"\{\s*id:'([^']*)'[^}]*phase:'([^']*)'[^}]*principle:'([^']*)'[^}]*"
        r"title:'((?:[^'\\]|\\.)*)'[^}]*desc:'((?:[^'\\]|\\.)*)'[^}]*"
        r"deps:\[([^\]]*)\]", re.DOTALL
    )
    for m in action_pattern.finditer(actions_match.group(1)):
        deps = [d.strip().strip("'") for d in m.group(6).split(',') if d.strip().strip("'")]
        actions_data.append({
            'id': m.group(1), 'phase': m.group(2), 'principle': m.group(3),
            'title': m.group(4).replace("\\'", "'"),
            'desc': m.group(5).replace("\\'", "'"), 'deps': deps,
        })
    print(f"Extracted {len(actions_data)} actions from ACTIONS")

# ── Parse HTML ──
soup = BeautifulSoup(html, 'html.parser')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def clean(text):
    if not text: return ''
    return re.sub(r'\s+', ' ', text).strip()

def grey(doc, text):
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = m.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(150, 150, 150)

def divider(doc, char='─', n=60):
    doc.add_paragraph(char * n)

# ══════════════════════════════════════════════════════════════
# DOCUMENT HEADER
# ══════════════════════════════════════════════════════════════
doc.add_heading('Nature Positive Roadmap — Complete Content Extract', level=0)
doc.add_paragraph(
    'This document contains ALL text content from the website including:\n'
    '• Visible page content (all sections)\n'
    '• Timeline entries with expandable details and stat pills\n'
    '• Challenge & Response cards with expandable details and stats\n'
    '• Roadmap grid (principles, targets, action chips, baselines)\n'
    '• 21 expandable target panels (context, actions, GBCA, industry)\n'
    '• 30 action network descriptions with dependencies\n'
    '• Stakeholder role tabs (Government, Developers, Manufacturers, Finance)\n'
    '• Case studies, TNFD steps, Green Star cards, GBCA pillars\n'
    '• Stats bar, image credits, footer content\n\n'
    'Section markers like [SECTION: ...] and [PANEL: ...] are for mapping back to HTML — do not remove them.'
)
divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 1: HERO SECTION
# ══════════════════════════════════════════════════════════════
hero = soup.find('section', id='hero') or soup.find('div', class_='hero')
if hero:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="hero" | class="hero"]')
    h1 = hero.find('h1')
    if h1:
        # Get main title and subtitle separately
        spans = h1.find_all('span')
        main_text = h1.get_text()
        for s in spans:
            main_text = main_text.replace(s.get_text(), '')
        doc.add_heading(clean(main_text), level=1)
        for s in spans:
            doc.add_paragraph(clean(s.get_text()))
    for p in hero.find_all('p'):
        t = clean(p.get_text())
        if t: doc.add_paragraph(t)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 2: STATS BAR
# ══════════════════════════════════════════════════════════════
stats = soup.find('div', class_='stats-section') or soup.find('section', class_='stats-section')
if stats:
    doc.add_paragraph('')
    grey(doc, '[SECTION: class="stats-section"]')
    doc.add_heading('Key Statistics', level=2)
    for item in stats.find_all('div', class_='stat-item'):
        num = item.find(class_='stat-number')
        label = item.find(class_='stat-label')
        if num and label:
            doc.add_paragraph(f'{clean(num.get_text())} — {clean(label.get_text())}')
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 3: ACKNOWLEDGEMENT OF COUNTRY
# ══════════════════════════════════════════════════════════════
ack = soup.find('div', class_='acknowledgement-section')
if ack:
    doc.add_paragraph('')
    grey(doc, '[SECTION: class="acknowledgement-section"]')
    doc.add_heading('Acknowledgement of Country', level=2)
    heading = ack.find(class_='ack-heading')
    if heading:
        doc.add_heading(clean(heading.get_text()), level=3)
    for p in ack.find_all('p'):
        if 'ack-heading' not in (p.get('class') or []):
            t = clean(p.get_text())
            if t: doc.add_paragraph(t)
    bq = ack.find('blockquote')
    if bq:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'"{clean(bq.get_text())}"')
        r.italic = True
    cite = ack.find('cite')
    if cite:
        p = doc.add_paragraph(f'— {clean(cite.get_text())}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 4: PURPOSE / FRAMEWORK
# ══════════════════════════════════════════════════════════════
purpose = soup.find('section', id='purpose') or soup.find('div', class_='purpose-section')
if purpose:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="purpose" | class="purpose-section"]')
    h2 = purpose.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    for p in purpose.find_all('p', class_=lambda c: c != 'purpose-form-label' if c else True):
        t = clean(p.get_text())
        if t and len(t) > 5:
            doc.add_paragraph(t)
    # Download form
    form = purpose.find('form') or purpose.find(class_='purpose-download')
    if form:
        doc.add_heading('Download the Roadmap', level=4)
        h4 = form.find('h4')
        if h4: doc.add_paragraph(clean(h4.get_text()))
        for p in form.find_all('p'):
            t = clean(p.get_text())
            if t: doc.add_paragraph(t)
        # Form fields
        for label in form.find_all(class_='purpose-form-label'):
            doc.add_paragraph(f'[Form field: {clean(label.get_text())}]')
        success = form.find(class_='purpose-form-success')
        if success:
            doc.add_paragraph(f'[Success message: {clean(success.get_text())}]')
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 5: POLICY LANDSCAPE / TIMELINE
# ══════════════════════════════════════════════════════════════
policy = soup.find('section', id='policy') or soup.find('section', class_='policy-section')
if policy:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="policy" | class="policy-section"]')
    h2 = policy.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    for p_tag in policy.find('div', class_='section-header').find_all('p') if policy.find('div', class_='section-header') else []:
        t = clean(p_tag.get_text())
        if t: doc.add_paragraph(t)

    # Column headings
    col_headings = policy.find(class_='tl-col-headings')
    if col_headings:
        global_h = col_headings.find(class_='tl-col-heading-global')
        aus_h = col_headings.find(class_='tl-col-heading-aus')
        if global_h and aus_h:
            doc.add_paragraph(f'[Column: {clean(global_h.get_text())} | {clean(aus_h.get_text())}]')

    # Timeline rows
    timeline = policy.find(class_='tl-timeline')
    if timeline:
        for row in timeline.find_all('div', class_='tl-row'):
            year_el = row.find(class_='tl-year')
            year = clean(year_el.get_text()) if year_el else ''
            is_milestone = 'tl-milestone' in (row.get('class') or [])

            entries = row.find_all(class_='tl-entry')
            for entry in entries:
                side = 'Global' if 'tl-global' in entry.get('class', []) else 'Australian'

                # Title (with tags)
                title_el = entry.find(class_='tl-entry-title')
                if not title_el: continue
                title = clean(title_el.get_text()).replace('▼', '').strip()

                # Extract tags
                tags = []
                for tag in title_el.find_all(class_='pt-tag'):
                    tags.append(clean(tag.get_text()))

                doc.add_paragraph('')
                milestone_marker = ' [MILESTONE]' if is_milestone else ''
                grey(doc, f'[TIMELINE: {year} | {side}{milestone_marker}]')

                tag_str = f' [{", ".join(tags)}]' if tags else ''
                doc.add_heading(f'{title}{tag_str}', level=3)

                # Expandable detail
                detail = entry.find(class_='tl-detail')
                if detail:
                    for p_tag in detail.find_all('p'):
                        t = clean(p_tag.get_text())
                        if t: doc.add_paragraph(t)
                    for li in detail.find_all('li'):
                        t = clean(li.get_text())
                        if t: doc.add_paragraph(t, style='List Bullet')
                    # Stat pills
                    for pill in detail.find_all(class_='pt-stat-pill'):
                        t = clean(pill.get_text())
                        if t: doc.add_paragraph(f'[Stat: {t}]')

    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 6: CHALLENGE & RESPONSE
# ══════════════════════════════════════════════════════════════
cr = soup.find('section', id='response') or soup.find('section', class_='cr-section')
if cr:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="response" | class="cr-section"]')
    h2 = cr.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    header_p = cr.find('div', class_='section-header')
    if header_p:
        for p_tag in header_p.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)

    # Column headers
    for gh in cr.find_all(class_='cr-grid-header'):
        t = clean(gh.get_text())
        if t: doc.add_paragraph(f'[Column: {t}]')

    # Challenge and response items - only top-level cr-item divs with data-pair
    for item in cr.find_all('div', attrs={'data-pair': True}):
        classes = item.get('class', [])

        pair = item.get('data-pair', '?')
        is_challenge = any('cr-item-ch' in c for c in classes)
        is_response = any('cr-item-pr' in c for c in classes)
        side = 'Challenge' if is_challenge else 'Response'

        h4 = item.find('h4')
        summary = item.find(class_='cr-item-summary')

        if h4:
            doc.add_paragraph('')
            grey(doc, f'[CR PAIR {pair}: {side}]')
            doc.add_heading(clean(h4.get_text()), level=3)

        if summary:
            doc.add_paragraph(clean(summary.get_text()))

        # Expandable detail
        detail = item.find(class_='cr-item-detail')
        if detail:
            for p_tag in detail.find_all('p'):
                t = clean(p_tag.get_text())
                if t: doc.add_paragraph(t)
            for li in detail.find_all('li'):
                t = clean(li.get_text())
                if t: doc.add_paragraph(t, style='List Bullet')
            # Stat pills
            for pill in detail.find_all(class_='cr-stat-pill'):
                t = clean(pill.get_text())
                if t: doc.add_paragraph(f'[Stat: {t}]')

    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 7: ROADMAP GRID
# ══════════════════════════════════════════════════════════════
roadmap = soup.find('section', id='roadmap')
if roadmap:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="roadmap" | class="irm-hero"]')
    h2 = roadmap.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    for p_tag in roadmap.find('div', class_='section-header').find_all('p') if roadmap.find('div', class_='section-header') else []:
        t = clean(p_tag.get_text())
        if t: doc.add_paragraph(t)

    subtitle = roadmap.find(class_='irm-hero-subtitle')
    if subtitle:
        doc.add_paragraph(clean(subtitle.get_text()))

    # Grid headers
    doc.add_heading('Roadmap Grid', level=2)
    for header in roadmap.find_all('div', class_='irm-header'):
        t = clean(header.get_text())
        if t and len(t) > 1:
            doc.add_paragraph(f'[Column: {t}]')

    # Principles, targets, action chips
    for principle in roadmap.find_all('div', class_='irm-principle'):
        t = clean(principle.get_text())
        if t:
            doc.add_heading(t, level=3)

    # All targets
    doc.add_heading('Targets by Principle', level=2)
    for target in roadmap.find_all('div', class_='irm-target'):
        t = clean(target.get_text()).replace('+', '').strip()
        classes = target.get('class', [])
        is_baseline = 't-baseline' in classes
        onclick = target.get('onclick', '')
        panel_id = ''
        if onclick:
            pm = re.search(r"openPanel\('([^']+)'\)", onclick)
            if pm: panel_id = f' → Panel: {pm.group(1)}'
        prefix = '[Baseline]' if is_baseline else '[Target]'
        doc.add_paragraph(f'{prefix} {t}{panel_id}', style='List Bullet')

    # All action chips
    doc.add_heading('Action Chips (Next 5 Years)', level=2)
    for chip in roadmap.find_all('div', class_='irm-action-chip'):
        t = clean(chip.get_text())
        if t: doc.add_paragraph(f'{t}', style='List Bullet')

    # Enabler row
    enabler = roadmap.find(class_='irm-enabler-row')
    if enabler:
        doc.add_paragraph(clean(enabler.get_text()))

    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 8: ACTION NETWORK SECTION
# ══════════════════════════════════════════════════════════════
actions_section = soup.find('section', id='actions')
if actions_section:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="actions" | class="action-net-section"]')
    h2 = actions_section.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    for p_tag in actions_section.find_all('p'):
        t = clean(p_tag.get_text())
        if t and len(t) > 10: doc.add_paragraph(t)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 9: PRINCIPLE DETAILS (with case studies + mini roadmaps)
# ══════════════════════════════════════════════════════════════
detail_div = soup.find('div', id='detail')
if detail_div:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="detail"]')
    doc.add_heading('Principle Details', level=1)

    for pd in detail_div.find_all('div', class_='principle-detail'):
        pnum = pd.find(class_='principle-num')
        h3 = pd.find('h3')
        if pnum:
            doc.add_paragraph(clean(pnum.get_text()))
        if h3:
            doc.add_heading(clean(h3.get_text()), level=2)

        # Main description
        for p_tag in pd.find_all('p', recursive=False):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
        text_div = pd.find(class_='principle-detail-text')
        if text_div:
            for p_tag in text_div.find_all('p', recursive=True):
                classes = p_tag.get('class', [])
                if 'image-credit' in classes:
                    continue
                parent_class = ' '.join(p_tag.parent.get('class', []))
                if 'case-study' in parent_class or 'mini-rm' in parent_class:
                    continue
                t = clean(p_tag.get_text())
                if t and len(t) > 10:
                    doc.add_paragraph(t)

        # Mini roadmap
        for mini in pd.find_all(class_='mini-roadmap'):
            label = mini.find(class_='mini-roadmap-label')
            if label:
                doc.add_heading(clean(label.get_text()), level=4)
            for card in mini.find_all(class_='mini-rm-card'):
                if card.get('style') and 'hidden' in card.get('style', ''):
                    continue
                t = clean(card.get_text())
                year_el = card.find_previous(class_='mini-rm-year')
                year = clean(year_el.get_text()) if year_el else ''
                is_baseline = 'baseline' in (card.get('class') or [])
                prefix = f'[{year}]' if year else ''
                tag = ' (Baseline)' if is_baseline else ''
                if t:
                    doc.add_paragraph(f'{prefix} {t}{tag}', style='List Bullet')

        # Case study
        for cs in pd.find_all(class_='case-study'):
            cs_h4 = cs.find('h4')
            if cs_h4:
                doc.add_heading(clean(cs_h4.get_text()), level=4)
            for p_tag in cs.find_all('p'):
                t = clean(p_tag.get_text())
                if t: doc.add_paragraph(t)

        # Image credit
        for credit in pd.find_all(class_='image-credit'):
            t = clean(credit.get_text())
            if t:
                p = doc.add_paragraph(f'[Image credit: {t}]')
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 10: ENABLERS
# ══════════════════════════════════════════════════════════════
enablers = soup.find('section', id='enablers')
if enablers:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="enablers" | class="enablers-section"]')
    h2 = enablers.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    header = enablers.find(class_='section-header')
    if header:
        for p_tag in header.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    for card in enablers.find_all(class_='enabler-card'):
        h3 = card.find('h3')
        if h3:
            doc.add_heading(clean(h3.get_text()), level=3)
        for p_tag in card.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 11: CASE STUDIES CAROUSEL
# ══════════════════════════════════════════════════════════════
cases = soup.find('section', id='casestudies')
if cases:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="casestudies" | class="case-carousel-section"]')
    h2 = cases.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    header = cases.find(class_='section-header')
    if header:
        for p_tag in header.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    for card in cases.find_all(class_='case-carousel-card'):
        h4 = card.find('h4')
        loc = card.find(class_='case-carousel-location')
        if h4:
            doc.add_heading(clean(h4.get_text()), level=3)
        if loc:
            doc.add_paragraph(clean(loc.get_text()))
        body_div = card.find(class_='case-carousel-body')
        if body_div:
            for p_tag in body_div.find_all('p'):
                t = clean(p_tag.get_text())
                if t: doc.add_paragraph(t)
        # Image alt
        img = card.find('img')
        if img and img.get('alt'):
            p = doc.add_paragraph(f'[Image: {img["alt"]}]')
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 12: TNFD
# ══════════════════════════════════════════════════════════════
tnfd = soup.find('section', id='tnfd')
if tnfd:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="tnfd" | class="tnfd-section"]')
    h2 = tnfd.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    header = tnfd.find(class_='section-header')
    if header:
        for p_tag in header.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    for step in tnfd.find_all(class_='tnfd-step'):
        num = step.find(class_='tnfd-step-num')
        h4 = step.find('h4')
        if num and h4:
            doc.add_heading(f'Step {clean(num.get_text())}: {clean(h4.get_text())}', level=3)
        elif h4:
            doc.add_heading(clean(h4.get_text()), level=3)
        for p_tag in step.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 13: GREEN STAR
# ══════════════════════════════════════════════════════════════
greenstar = soup.find('section', id='greenstar')
if greenstar:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="greenstar" | class="greenstar-section"]')
    h2 = greenstar.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    header = greenstar.find(class_='section-header')
    if header:
        for p_tag in header.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    for card in greenstar.find_all(class_='gs-card'):
        h3 = card.find('h3')
        if h3:
            doc.add_heading(clean(h3.get_text()), level=3)
        for p_tag in card.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
        for li in card.find_all('li'):
            t = clean(li.get_text())
            if t: doc.add_paragraph(t, style='List Bullet')
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 14: GBCA ROLE
# ══════════════════════════════════════════════════════════════
gbca_role = soup.find('section', class_='gbca-role-section')
if not gbca_role:
    # Try finding by content
    for s in soup.find_all('section'):
        if s.find('div', class_='gbca-pillars'):
            gbca_role = s
            break
if gbca_role:
    doc.add_paragraph('')
    grey(doc, '[SECTION: class="gbca-role-section"]')
    h2 = gbca_role.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    header = gbca_role.find(class_='section-header')
    if header:
        for p_tag in header.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    for card in gbca_role.find_all(class_='pillar-card'):
        num = card.find(class_='pillar-num')
        h4 = card.find('h4')
        if h4:
            prefix = f'{clean(num.get_text())}. ' if num else ''
            doc.add_heading(f'{prefix}{clean(h4.get_text())}', level=3)
        for p_tag in card.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 15: STAKEHOLDER ROLES (TABBED)
# ══════════════════════════════════════════════════════════════
action_section = soup.find('section', id='action')
if action_section:
    doc.add_paragraph('')
    grey(doc, '[SECTION: id="action" | Stakeholder Roles — Tabbed Content]')
    h2 = action_section.find('h2')
    if h2:
        doc.add_heading(clean(h2.get_text()), level=1)
    header = action_section.find(class_='section-header')
    if header:
        for p_tag in header.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)

    tab_names = {'tab-gov': 'Government', 'tab-dev': 'Developers', 'tab-mfg': 'Manufacturers', 'tab-fin': 'Finance Sector'}
    for tab_id, tab_name in tab_names.items():
        panel = action_section.find(id=tab_id)
        if not panel: continue

        doc.add_paragraph('')
        grey(doc, f'[TAB: {tab_id}]')
        doc.add_heading(f'Tab: {tab_name}', level=2)

        for card in panel.find_all(class_='role-card'):
            h3 = card.find('h3')
            if h3:
                doc.add_heading(clean(h3.get_text()), level=3)
            for p_tag in card.find_all('p'):
                t = clean(p_tag.get_text())
                if t: doc.add_paragraph(t)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 16: CLOSING QUOTE
# ══════════════════════════════════════════════════════════════
quote_blocks = soup.find_all('div', class_='quote-block')
for qb in quote_blocks:
    doc.add_paragraph('')
    grey(doc, '[SECTION: class="quote-block"]')
    bq = qb.find('blockquote')
    if bq:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'"{clean(bq.get_text())}"')
        r.italic = True
    cite = qb.find('cite')
    if cite:
        p = doc.add_paragraph(f'— {clean(cite.get_text())}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 17: FOOTER
# ══════════════════════════════════════════════════════════════
footer = soup.find('footer')
if footer:
    doc.add_paragraph('')
    grey(doc, '[SECTION: <footer>]')
    doc.add_heading('Footer', level=2)
    brand = footer.find(class_='footer-brand')
    if brand:
        h3 = brand.find('h3')
        if h3: doc.add_heading(clean(h3.get_text()), level=3)
        for p_tag in brand.find_all('p'):
            t = clean(p_tag.get_text())
            if t: doc.add_paragraph(t)
    links = footer.find(class_='footer-links')
    if links:
        h4 = links.find('h4')
        if h4: doc.add_heading(clean(h4.get_text()), level=4)
        for a in links.find_all('a'):
            t = clean(a.get_text())
            if t: doc.add_paragraph(t, style='List Bullet')
    bottom = footer.find(class_='footer-bottom')
    if bottom:
        t = clean(bottom.get_text())
        if t: doc.add_paragraph(t)
    divider(doc)

# ══════════════════════════════════════════════════════════════
# PART 18: INTERACTIVE ROADMAP PANEL DATA (JS)
# ══════════════════════════════════════════════════════════════
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_heading('INTERACTIVE ROADMAP — Expandable Panel Details', level=1)
doc.add_paragraph(
    'The following content appears when users click the "+" buttons on roadmap targets. '
    'Each panel contains context, required actions, what GBCA is doing, and what industry needs to do.'
)
divider(doc)

principle_order = [
    ('Principle 1: Prevent Nature Loss', ['p1-measure', 'p1-nonetloss', 'p1-noloss']),
    ('Principle 2: Increase & Connect Nature', ['p2-measure', 'p2-bng', 'p2-additional']),
    ('Principle 3: Drive Circularity', ['p3-sector-measure', 'p3-sector-7', 'p3-sector-15', 'p3-sector-future',
                                         'p3-proj-measure', 'p3-proj-10', 'p3-proj-20', 'p3-proj-future']),
    ('Principle 4: Choose Low-Impact Materials', ['p4-measure', 'p4-implement', 'p4-bng', 'p4-eliminate']),
    ('Principle 5: Invest in Nature', ['p5-measure', 'p5-invest', 'p5-commonplace']),
]

for principle_name, panel_ids in principle_order:
    doc.add_heading(principle_name, level=2)
    for pid in panel_ids:
        if pid not in panel_data: continue
        p = panel_data[pid]
        doc.add_paragraph('')
        grey(doc, f'[PANEL: {pid}]')
        doc.add_heading(p['title'], level=3)
        doc.add_paragraph(f"({p['time']})")
        doc.add_heading('Context', level=4)
        doc.add_paragraph(p['context'])
        doc.add_heading('Actions Needed in the Next Five Years', level=4)
        for item in p['actions']:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_heading('What GBCA Is Doing', level=4)
        for item in p['gbca']:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_heading('What Industry Needs to Do', level=4)
        for item in p['industry']:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph('─' * 40)

# ══════════════════════════════════════════════════════════════
# PART 19: ACTION NETWORK (30 actions from JS)
# ══════════════════════════════════════════════════════════════
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_heading('ACTION NETWORK — 30 Actions (shown on hover)', level=1)
doc.add_paragraph(
    'The following 30 actions appear in the interactive action network diagram. '
    'Each has a title, description, phase, principle, and dependencies.'
)
divider(doc)

principle_map = {
    'prevent': 'Principle 1: Prevent Nature Loss',
    'increase': 'Principle 2: Increase & Connect Nature',
    'circular': 'Principle 3: Drive Circularity',
    'materials': 'Principle 4: Choose Low-Impact Materials',
    'invest': 'Principle 5: Invest in Nature',
}
phase_map = {'define': 'Define', 'advocate': 'Advocate', 'implement': 'Implement'}
action_lookup = {a['id']: a['title'] for a in actions_data}

for principle_id, principle_label in principle_map.items():
    principle_actions = [a for a in actions_data if a['principle'] == principle_id]
    if not principle_actions: continue
    doc.add_heading(principle_label, level=2)
    for a in principle_actions:
        doc.add_paragraph('')
        grey(doc, f'[ACTION: {a["id"]}]')
        doc.add_heading(f'{a["id"].upper()}: {a["title"]}', level=3)
        doc.add_paragraph(f'Phase: {phase_map.get(a["phase"], a["phase"])}')
        doc.add_paragraph(a['desc'])
        if a['deps']:
            dep_names = [f'{d} ({action_lookup.get(d, "?")})' for d in a['deps']]
            doc.add_paragraph(f'Depends on: {", ".join(dep_names)}')

divider(doc)

# Save
output_path = '/home/user/nature-roadmap-v3/Nature_Positive_Roadmap_Content.docx'
doc.save(output_path)
print(f"\nDone! Saved to {output_path}")

# Count content
total_p = len(doc.paragraphs)
headings = len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])
bullets = len([p for p in doc.paragraphs if p.style.name == 'List Bullet'])
print(f"Total paragraphs: {total_p}")
print(f"Headings: {headings}")
print(f"Bullet points: {bullets}")
print(f"Panels: {len(panel_data)}")
print(f"Actions: {len(actions_data)}")
